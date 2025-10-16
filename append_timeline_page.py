#!/usr/bin/env python3
"""Build a timeline infographic and append it to an existing report."""

import argparse
import json
import math
import os
import tempfile
from collections import defaultdict
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import FancyBboxPatch
from matplotlib.ticker import MaxNLocator
from PIL import Image


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--events_json",
        help="Path to events JSON or JSON string (if omitted, demo data is generated)",
    )
    parser.add_argument(
        "--focus_scores_json",
        help="Optional path or JSON string for focus scores",
    )
    parser.add_argument("--input_report", help="Existing report to append to", nargs="?")
    parser.add_argument("--output_report", help="Output report path", nargs="?")
    parser.add_argument(
        "--title",
        default="Year-at-a-Glance Timeline",
        help="Title for the infographic",
    )
    parser.add_argument(
        "--palette_json",
        help="Optional palette overrides as JSON string or path",
    )
    return parser.parse_args()


def load_json_input(maybe_path_or_json: Optional[str]) -> Optional[Dict[str, Any]]:
    if not maybe_path_or_json:
        return None
    if os.path.exists(maybe_path_or_json):
        with open(maybe_path_or_json, "r", encoding="utf-8") as handle:
            return json.load(handle)
    try:
        return json.loads(maybe_path_or_json)
    except json.JSONDecodeError:
        raise ValueError(f"Unable to parse JSON from '{maybe_path_or_json}'.")


def synthesize_demo_events(year: int = 2024) -> Dict[str, Any]:
    events = [
        {"date": f"{year}-01-18", "transit_body": "Sun", "natal_body": "Midheaven", "aspect": "conjunction", "orb": 1.2, "score": 85},
        {"date": f"{year}-03-09", "transit_body": "Venus", "natal_body": "Moon", "aspect": "trine", "orb": 0.5, "score": 65},
        {"date": f"{year}-04-21", "transit_body": "Mars", "natal_body": "Saturn", "aspect": "square", "orb": 1.1, "score": -70},
        {"date": f"{year}-06-02", "transit_body": "Jupiter", "natal_body": "Venus", "aspect": "sextile", "orb": 0.4, "score": 55},
        {"date": f"{year}-08-14", "transit_body": "Saturn", "natal_body": "Sun", "aspect": "opposition", "orb": 1.6, "score": -80},
        {"date": f"{year}-10-05", "transit_body": "Mercury", "natal_body": "Ascendant", "aspect": "trine", "orb": 0.8, "score": 45},
        {"date": f"{year}-11-22", "transit_body": "Venus", "natal_body": "Mars", "aspect": "conjunction", "orb": 0.3, "score": 60},
        {"date": f"{year}-12-28", "transit_body": "Sun", "natal_body": "TrueNode", "aspect": "square", "orb": 1.0, "score": -50},
    ]
    months: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for event in events:
        month_key = event["date"][:7]
        months[month_key].append(event)
    return {
        "meta": {"year": year},
        "months": months,
        "top_events": events[:4],
    }


def synthesize_focus_scores(year: int) -> Dict[str, Any]:
    sections = [
        "Career",
        "Money",
        "Love",
        "Health",
        "Growth",
        "Home/Family",
        "Social",
        "Spiritual",
    ]
    months = {}
    rng = np.random.default_rng(seed=year)
    for month in range(1, 13):
        base = 50 + 15 * np.sin((month / 12) * 2 * np.pi)
        month_scores = {section: int(np.clip(base + rng.normal(0, 10), 35, 85)) for section in sections}
        months[f"{year}-{month:02d}"] = month_scores
    return {"year": year, "sections": sections, "months": months}


def ensure_output_dir(path: str) -> None:
    directory = os.path.dirname(os.path.abspath(path))
    if directory and not os.path.exists(directory):
        os.makedirs(directory, exist_ok=True)


def build_palette(palette_json: Optional[str]) -> Dict[str, Any]:
    palette = {
        "bg": "#0f172a",
        "grid": "#1e293b",
        "text": "#f8fafc",
        "positive": "#22c55e",
        "negative": "#ef4444",
        "timeline": "#94a3b8",
        "sections": [
            "#2563eb",
            "#10b981",
            "#f59e0b",
            "#ef4444",
            "#8b5cf6",
            "#059669",
            "#dc2626",
            "#0ea5e9",
        ],
    }
    if palette_json:
        loaded = load_json_input(palette_json)
        if loaded:
            palette.update({k: v for k, v in loaded.items() if k != "sections"})
            if "sections" in loaded:
                palette["sections"] = loaded["sections"]
    return palette


def flatten_events(events_json: Dict[str, Any]) -> Tuple[List[Dict[str, Any]], int]:
    if not events_json:
        events_json = synthesize_demo_events()
    year = events_json.get("meta", {}).get("year") or datetime.now().year
    all_events: List[Dict[str, Any]] = []
    for month_events in events_json.get("months", {}).values():
        all_events.extend(month_events)
    if not all_events:
        all_events = synthesize_demo_events(year)["top_events"]
    return all_events, year


ASPECT_POLARITY = {
    "conjunction": 0.6,
    "trine": 1.0,
    "sextile": 0.7,
    "square": -0.8,
    "opposition": -1.0,
}


SECTION_PLANETS = {
    "Career": {"Saturn", "Midheaven"},
    "Money": {"Jupiter", "Venus"},
    "Love": {"Venus", "Mars", "Moon"},
    "Health": {"Sun", "Mars"},
    "Growth": {"Jupiter", "Chiron"},
    "Home/Family": {"Moon", "Ascendant", "IC"},
    "Social": {"Mercury", "Venus"},
    "Spiritual": {"Neptune", "Pluto", "TrueNode"},
}


def event_polarity(event: Dict[str, Any]) -> float:
    aspect = str(event.get("aspect", "")).lower()
    if aspect in ASPECT_POLARITY:
        base = ASPECT_POLARITY[aspect]
    else:
        score = event.get("score", 0)
        base = math.copysign(1.0, score) if score else 0.0
    return base


def derive_focus_scores(
    focus_json: Optional[Dict[str, Any]],
    events: List[Dict[str, Any]],
    year: int,
) -> Dict[str, Dict[str, float]]:
    month_keys = [f"{year}-{m:02d}" for m in range(1, 13)]
    provided_focus: Dict[str, Dict[str, float]] = {}
    if focus_json and focus_json.get("months"):
        provided = focus_json["months"]
        for month in month_keys:
            if month in provided:
                provided_focus[month] = {k: float(v) for k, v in provided[month].items()}
        if len(provided_focus) == 12:
            return provided_focus
    # Derive synthetic focus scores from events
    month_section_scores: Dict[str, Dict[str, float]] = {
        month: {section: 0.0 for section in SECTION_PLANETS} for month in month_keys
    }
    month_counts: Dict[str, Dict[str, int]] = {
        month: {section: 0 for section in SECTION_PLANETS} for month in month_keys
    }
    for event in events:
        date_str = event.get("date")
        if not date_str:
            continue
        month_key = date_str[:7]
        if month_key not in month_section_scores:
            continue
        transit = event.get("transit_body", "")
        natal = event.get("natal_body", "")
        sign = event_polarity(event)
        raw_score = float(event.get("score", 0))
        signed_score = raw_score * (1 if sign >= 0 else -1)
        for section, bodies in SECTION_PLANETS.items():
            if transit in bodies or natal in bodies:
                month_section_scores[month_key][section] += signed_score
                month_counts[month_key][section] += 1
    normalized: Dict[str, Dict[str, float]] = {month: {} for month in month_keys}
    for section in SECTION_PLANETS:
        values = [month_section_scores[m][section] for m in month_keys]
        if max(values) == min(values):
            norm_values = [50.0 for _ in values]
        else:
            min_v = min(values)
            max_v = max(values)
            norm_values = [float(np.clip((v - min_v) / (max_v - min_v) * 100.0, 0, 100)) for v in values]
        last_val = 50.0
        for idx, month in enumerate(month_keys):
            val = norm_values[idx]
            if month_counts[month][section] == 0:
                val = last_val
            normalized[month][section] = val
            last_val = val
    for month, provided_values in provided_focus.items():
        normalized.setdefault(month, {}).update(provided_values)
    return normalized


def prepare_highlights(events: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    positives = []
    negatives = []
    for event in events:
        sign = event_polarity(event)
        if sign >= 0:
            positives.append(event)
        if sign < 0:
            negatives.append(event)
    positives.sort(key=lambda e: abs(e.get("score", 0)), reverse=True)
    negatives.sort(key=lambda e: abs(e.get("score", 0)), reverse=True)
    positives = sorted(positives[:8], key=lambda e: e.get("date", ""))
    negatives = sorted(negatives[:8], key=lambda e: e.get("date", ""))
    return positives, negatives


def format_event_label(event: Dict[str, Any]) -> str:
    date_str = event.get("date", "")
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        date_fmt = date_obj.strftime("%b %d")
    except Exception:
        date_fmt = date_str
    parts = [date_fmt]
    transit = event.get("transit_body")
    natal = event.get("natal_body")
    if transit or natal:
        parts.append("{}{}{}".format(transit or "", "–" if transit and natal else "", natal or ""))
    aspect = event.get("aspect")
    if aspect:
        parts.append(str(aspect))
    score = event.get("score")
    if score is not None:
        parts.append(f"{score:+.0f}")
    return " • ".join(parts)


def node_radius(score: float, min_r: float = 4.0, max_r: float = 14.0) -> float:
    magnitude = abs(score)
    return min_r + (max_r - min_r) * min(magnitude, 100) / 100.0


def timeline_position(date_str: str, year: int, width: float, margin: float) -> float:
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
    except Exception:
        return margin
    start = datetime(year, 1, 1)
    end = datetime(year + 1, 1, 1)
    total_days = (end - start).days
    delta_days = (date_obj - start).days
    fraction = np.clip(delta_days / total_days, 0, 1)
    return margin + fraction * (width - 2 * margin)


def draw_highlight_row(
    ax: matplotlib.axes.Axes,
    events: List[Dict[str, Any]],
    start_x: float,
    center_y: float,
    color: str,
    palette: Dict[str, Any],
    total_width: float,
    spacing: float = 20.0,
) -> None:
    if not events:
        return
    count = len(events)
    pill_height = 42.0
    pill_width = (total_width - spacing * (count - 1)) / count if count > 0 else total_width
    pill_width = max(110.0, pill_width)
    for idx, event in enumerate(events):
        x = start_x + idx * (pill_width + spacing)
        patch = FancyBboxPatch(
            (x, center_y - pill_height / 2),
            pill_width,
            pill_height,
            boxstyle="round,pad=0.6",
            linewidth=0,
            facecolor=color,
            alpha=0.9,
        )
        ax.add_patch(patch)
        ax.text(
            x + pill_width / 2,
            center_y,
            format_event_label(event),
            ha="center",
            va="center",
            color="#f8fafc",
            fontsize=14,
            fontweight="bold",
        )


def draw_timeline(
    ax: matplotlib.axes.Axes,
    events: List[Dict[str, Any]],
    palette: Dict[str, Any],
    width: float,
    height: float,
    margin: float,
    year: int,
    focus_top: float,
) -> None:
    top_area_height = height - focus_top - margin
    timeline_y = focus_top + top_area_height * 0.55
    line_color = palette["timeline"]
    ax.plot([margin, width - margin], [timeline_y, timeline_y], color=line_color, linewidth=3)
    months = [datetime(year, m, 1) for m in range(1, 13)]
    for month in months:
        x = timeline_position(month.strftime("%Y-%m-%d"), year, width, margin)
        ax.plot([x, x], [timeline_y - 12, timeline_y + 12], color=line_color, linewidth=1.5)
        ax.text(
            x,
            timeline_y + 22,
            month.strftime("%b").upper(),
            ha="center",
            va="bottom",
            color=palette["text"],
            fontsize=14,
            fontweight="bold",
        )
    positive_offsets = [80, 130, 180]
    negative_offsets = [80, 130, 180]
    pos_index = 0
    neg_index = 0
    last_x_side: Dict[str, float] = {"positive": -9999.0, "negative": -9999.0}
    for idx, event in enumerate(sorted(events, key=lambda e: e.get("date", ""))):
        date_str = event.get("date", "")
        x = timeline_position(date_str, year, width, margin)
        score = float(event.get("score", 0))
        sign = event_polarity(event)
        color = palette["positive"] if sign >= 0 else palette["negative"] if sign < 0 else palette["timeline"]
        radius = node_radius(score)
        connector = None
        if sign >= 0.05:
            offset = positive_offsets[pos_index % len(positive_offsets)]
            pos_index += 1
            label_y = timeline_y + offset + 20
            node_y = timeline_y + offset
            connector = [timeline_y + 12, node_y - radius]
            side = "positive"
            va = "bottom"
        elif sign <= -0.05:
            offset = negative_offsets[neg_index % len(negative_offsets)]
            neg_index += 1
            label_y = timeline_y - offset - 20
            node_y = timeline_y - offset
            connector = [timeline_y - 12, node_y + radius]
            side = "negative"
            va = "top"
        else:
            node_y = timeline_y
            label_y = timeline_y + 40
            side = "neutral"
            va = "bottom"
        if side in last_x_side and abs(x - last_x_side[side]) < 45:
            if side == "positive":
                node_y += 18
                label_y += 18
            elif side == "negative":
                node_y -= 18
                label_y -= 18
        if side in last_x_side:
            last_x_side[side] = x
        if connector and connector[0] != connector[1]:
            ax.plot([x, x], connector, color=color, linewidth=1.2)
        circle = plt.Circle((x, node_y), radius, color=color, ec="white", lw=1.5)
        ax.add_patch(circle)
        rotation = ((idx % 5) - 2) * 4
        rotation = float(np.clip(rotation, -12, 12))
        ax.text(
            x,
            label_y,
            format_event_label(event),
            ha="center",
            va=va,
            rotation=rotation,
            rotation_mode="anchor",
            color=palette["text"],
            fontsize=13,
        )


def draw_focus_panels(
    fig: matplotlib.figure.Figure,
    focus_scores: Dict[str, Dict[str, float]],
    palette: Dict[str, Any],
    width: float,
    height: float,
    margin: float,
    focus_height: float,
) -> None:
    rows, cols = 2, 4
    gap_x = 20.0
    gap_y = 24.0
    panel_width = (width - 2 * margin - (cols - 1) * gap_x) / cols
    panel_height = (focus_height - (rows - 1) * gap_y) / rows
    focus_bottom = margin
    month_keys = sorted(focus_scores.keys())
    sections = list(SECTION_PLANETS.keys())
    for idx, section in enumerate(sections):
        row = idx // cols
        col = idx % cols
        left = margin + col * (panel_width + gap_x)
        bottom = focus_bottom + row * (panel_height + gap_y)
        ax = fig.add_axes([
            left / width,
            bottom / height,
            panel_width / width,
            panel_height / height,
        ])
        values = [focus_scores.get(month, {}).get(section, 50.0) for month in month_keys]
        if len(values) < 12:
            filler = values[-1] if values else 50.0
            values = values + [filler] * (12 - len(values))
        values = values[:12]
        avg = float(np.nanmean(values)) if values else 0.0
        ax.set_facecolor(palette["bg"])
        ax.set_xlim(0, 11)
        ax.set_ylim(0, 100)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_visible(False)
        ax.spines["bottom"].set_visible(False)
        ax.tick_params(axis="both", which="both", length=0, labelsize=10, colors=palette["text"])
        ax.yaxis.set_major_locator(MaxNLocator(3))
        ax.set_xticks([0, 3, 6, 9, 11])
        ax.set_xticklabels(["Jan", "Apr", "Jul", "Oct", "Dec"], color=palette["text"], fontsize=9)
        ax.grid(color=palette["grid"], linestyle="-", linewidth=0.8, alpha=0.6)
        color = palette["sections"][idx % len(palette["sections"])]
        ax.plot(range(12), values, color=color, linewidth=2.5, marker="o", markersize=4, alpha=0.9)
        ax.fill_between(range(12), values, 50, color=color, alpha=0.08)
        ax.set_title(
            f"{section} • avg {avg:.0f}",
            fontsize=12,
            color=palette["text"],
            pad=6,
        )


def render_infographic(
    events: List[Dict[str, Any]],
    focus_scores: Dict[str, Dict[str, float]],
    title: str,
    palette: Dict[str, Any],
    year: int,
    output_path: Optional[str] = None,
) -> str:
    width, height = 1920.0, 1080.0
    margin = 60.0
    focus_height = height * 0.35
    matplotlib.rcParams["font.family"] = "DejaVu Sans"
    fig = plt.figure(figsize=(width / 100, height / 100), dpi=100)
    fig.patch.set_facecolor(palette["bg"])
    ax = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, width)
    ax.set_ylim(0, height)
    ax.axis("off")

    # Title
    ax.text(
        margin,
        height - margin,
        title,
        ha="left",
        va="top",
        color=palette["text"],
        fontsize=32,
        fontweight="bold",
    )
    ax.text(
        margin,
        height - margin - 36,
        f"{year} Year-at-a-Glance",
        ha="left",
        va="top",
        color=palette["text"],
        fontsize=18,
        alpha=0.85,
    )

    # Legend
    legend_x = width - margin - 240
    legend_y = height - margin
    legend_items = [
        (palette["positive"], "Positive"),
        (palette["negative"], "Challenging"),
        (palette["timeline"], "Neutral / Timeline"),
    ]
    ax.text(
        legend_x,
        legend_y,
        "Legend",
        ha="left",
        va="top",
        color=palette["text"],
        fontsize=18,
        fontweight="bold",
    )
    for idx, (color, label) in enumerate(legend_items):
        y = legend_y - 24 - idx * 26
        ax.add_patch(FancyBboxPatch((legend_x, y - 12), 20, 20, boxstyle="round,pad=0.2", facecolor=color, linewidth=0))
        ax.text(legend_x + 28, y - 2, label, ha="left", va="center", color=palette["text"], fontsize=13)
    ax.text(
        legend_x,
        legend_y - 24 - len(legend_items) * 26 - 10,
        "Dot size = event strength",
        ha="left",
        va="top",
        color=palette["text"],
        fontsize=12,
        alpha=0.8,
    )

    # Highlights
    positive_highlights, negative_highlights = prepare_highlights(events)
    highlight_width = width * 0.45
    highlight_left = margin
    highlight_width = width - 2 * margin
    draw_highlight_row(
        ax,
        positive_highlights,
        highlight_left,
        height - margin - 90,
        palette["positive"],
        palette,
        highlight_width,
    )
    draw_highlight_row(
        ax,
        negative_highlights,
        highlight_left,
        height - margin - 150,
        palette["negative"],
        palette,
        highlight_width,
    )

    # Timeline
    focus_top = margin + focus_height
    draw_timeline(ax, events, palette, width, height, margin, year, focus_top)

    # Focus panels
    draw_focus_panels(fig, focus_scores, palette, width, height, margin, focus_height)

    output_path = os.path.abspath(output_path or "yearly_forecast_timeline.png")
    fig.savefig(output_path, dpi=200, facecolor=palette["bg"], bbox_inches="tight", pad_inches=0)
    plt.close(fig)
    return output_path


def append_to_docx(input_path: str, output_path: str, png_path: str) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document(input_path)
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for container in (section.header, section.footer):
        for paragraph in container.paragraphs:
            paragraph.text = ""
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Year-at-a-Glance Timeline")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = "DejaVu Sans"
    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    max_width = (section.page_width - section.left_margin - section.right_margin) / 914400
    picture_paragraph.add_run().add_picture(png_path, width=Inches(max_width))
    ensure_output_dir(output_path)
    document.save(output_path)


def append_to_pdf(input_path: str, output_path: str, png_path: str) -> None:
    from pypdf import PdfReader, PdfWriter
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.pdfgen import canvas as rl_canvas

    reader = PdfReader(input_path)
    if reader.pages:
        first_page = reader.pages[0]
        width = float(first_page.mediabox.width)
        height = float(first_page.mediabox.height)
    else:
        width, height = landscape(letter)
    if width < height:
        width, height = landscape(letter)
    tmp_fd, tmp_pdf = tempfile.mkstemp(suffix="_timeline_tmp.pdf")
    os.close(tmp_fd)
    with Image.open(png_path) as img:
        img_width, img_height = img.size
    scale = min(width / img_width, height / img_height)
    draw_width = img_width * scale
    draw_height = img_height * scale
    pos_x = (width - draw_width) / 2
    pos_y = (height - draw_height) / 2
    c = rl_canvas.Canvas(tmp_pdf, pagesize=(width, height))
    c.drawImage(png_path, pos_x, pos_y, width=draw_width, height=draw_height, preserveAspectRatio=True)
    c.showPage()
    c.save()
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    timeline_reader = PdfReader(tmp_pdf)
    for page in timeline_reader.pages:
        writer.add_page(page)
    ensure_output_dir(output_path)
    with open(output_path, "wb") as handle:
        writer.write(handle)
    os.remove(tmp_pdf)


def export_pdf_only(output_path: str, png_path: str) -> None:
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.pdfgen import canvas as rl_canvas

    width, height = landscape(letter)
    with Image.open(png_path) as img:
        img_width, img_height = img.size
    scale = min(width / img_width, height / img_height)
    draw_width = img_width * scale
    draw_height = img_height * scale
    pos_x = (width - draw_width) / 2
    pos_y = (height - draw_height) / 2
    ensure_output_dir(output_path)
    c = rl_canvas.Canvas(output_path, pagesize=(width, height))
    c.drawImage(png_path, pos_x, pos_y, width=draw_width, height=draw_height, preserveAspectRatio=True)
    c.showPage()
    c.save()


def main() -> None:
    args = parse_args()
    events_json = load_json_input(args.events_json) if args.events_json else None
    events, year = flatten_events(events_json)
    focus_json = load_json_input(args.focus_scores_json) if args.focus_scores_json else None
    if focus_json is None and args.events_json is None:
        focus_json = synthesize_focus_scores(year)
    focus_scores = derive_focus_scores(focus_json, events, year)
    palette = build_palette(args.palette_json)
    png_path = render_infographic(events, focus_scores, args.title, palette, year)

    if args.input_report:
        if not args.output_report:
            raise ValueError("When --input_report is provided, --output_report must also be supplied.")
    if args.input_report and args.output_report:
        input_ext = os.path.splitext(args.input_report)[1].lower()
        output_ext = os.path.splitext(args.output_report)[1].lower()
        if input_ext == ".docx" and output_ext == ".docx":
            append_to_docx(args.input_report, args.output_report, png_path)
        elif input_ext == ".pdf" and output_ext == ".pdf":
            append_to_pdf(args.input_report, args.output_report, png_path)
        else:
            raise ValueError("Input and output report types must match and be either DOCX or PDF.")
    else:
        export_pdf_only(os.path.abspath("timeline_page.pdf"), png_path)


if __name__ == "__main__":
    main()

